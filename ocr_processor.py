"""OCR and image processing module for SmartDoc AI."""

from __future__ import annotations

import os
import re
import tempfile
import unicodedata
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass

import cv2
import numpy as np
import pypdfium2 as pdfium
from rapidocr_onnxruntime import RapidOCR
from docx import Document as DocxDocument
from langchain_core.documents import Document

try:
    import easyocr  # type: ignore
except Exception:  # pragma: no cover - optional dependency fallback
    easyocr = None


@dataclass
class OCRConfig:
    """Configuration for OCR processing."""
    # Supported languages for OCR
    languages: List[str] = None  # Default: Vietnamese, English, Chinese
    gpu: bool = False
    model_storage_directory: Optional[str] = None
    confidence_threshold: float = 0.3
    
    def __post_init__(self):
        if self.languages is None:
            self.languages = ['vi', 'en']  # Vietnamese and English are safe together by default


class OCRProcessor:
    """Process OCR for images extracted from PDF/DOCX files."""
    
    _reader_cache: Optional[RapidOCR] = None
    _easy_reader_cache: Optional[Any] = None
    _current_lang_config: Optional[OCRConfig] = None
    
    @classmethod
    def _get_reader(cls, config: OCRConfig) -> RapidOCR:
        """Get or create RapidOCR reader (cached)."""
        if cls._reader_cache is not None:
            return cls._reader_cache

        print("[*] Initializing OCR with RapidOCR")
        reader = RapidOCR()
        cls._reader_cache = reader
        cls._current_lang_config = OCRConfig(
            languages=list(config.languages or ["vi", "en"]),
            gpu=config.gpu,
            model_storage_directory=config.model_storage_directory,
            confidence_threshold=config.confidence_threshold,
        )
        return reader

    @classmethod
    def _get_easy_reader(cls, config: OCRConfig) -> Optional[Any]:
        """Get optional EasyOCR reader for hard cases like low-contrast stamps."""
        if easyocr is None:
            return None
        if cls._easy_reader_cache is not None:
            return cls._easy_reader_cache

        languages = [lang for lang in (config.languages or ["vi", "en"]) if lang]
        if not languages:
            languages = ["vi", "en"]

        try:
            print("[*] Initializing OCR fallback with EasyOCR")
            cls._easy_reader_cache = easyocr.Reader(languages, gpu=bool(config.gpu), verbose=False)
        except Exception as e:
            print(f"[!] EasyOCR init failed, keep using RapidOCR only: {e}")
            cls._easy_reader_cache = None

        return cls._easy_reader_cache

    @staticmethod
    def _enhance_image(image: np.ndarray) -> np.ndarray:
        """Enhance image quality for better OCR results."""
        if image is None or image.size == 0:
            return image

        work_image = image
        if len(work_image.shape) == 3:
            work_image = cv2.cvtColor(work_image, cv2.COLOR_BGR2GRAY)

        height, width = work_image.shape[:2]
        min_edge = min(height, width)
        
        # Downscale very large images to avoid OCR slowdown
        if min_edge > 1200:
            scale = 1200 / min_edge
            new_width = max(1, int(width * scale))
            new_height = max(1, int(height * scale))
            work_image = cv2.resize(work_image, (new_width, new_height), interpolation=cv2.INTER_AREA)
        elif min_edge < 300:
            scale = 300 / max(min_edge, 1)
            work_image = cv2.resize(
                work_image,
                (max(1, int(width * scale)), max(1, int(height * scale))),
                interpolation=cv2.INTER_CUBIC,
            )

        clahe = cv2.createCLAHE(clipLimit=2.5, tileGridSize=(8, 8))
        work_image = clahe.apply(work_image)
        work_image = cv2.fastNlMeansDenoising(work_image, None, 9, 7, 21)

        # Sharpen text edges, especially useful for faint stamp borders.
        kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]], dtype=np.float32)
        work_image = cv2.filter2D(work_image, -1, kernel)

        return work_image

    @staticmethod
    def _crop_center(image: np.ndarray, ratio: float = 0.72) -> np.ndarray:
        """Crop the center area to prioritize the main stamp text."""
        if image is None or image.size == 0:
            return image

        height, width = image.shape[:2]
        crop_height = max(1, int(height * ratio))
        crop_width = max(1, int(width * ratio))
        top = max(0, (height - crop_height) // 2)
        left = max(0, (width - crop_width) // 2)
        return image[top:top + crop_height, left:left + crop_width]

    @staticmethod
    def _build_red_stamp_variants(image: np.ndarray) -> List[np.ndarray]:
        """Build variants optimized for red stamps and seals."""
        if image is None or image.size == 0:
            return []

        work_image = image
        if len(work_image.shape) == 2:
            work_image = cv2.cvtColor(work_image, cv2.COLOR_GRAY2BGR)

        hsv = cv2.cvtColor(work_image, cv2.COLOR_BGR2HSV)
        lower_red_1 = np.array([0, 60, 50], dtype=np.uint8)
        upper_red_1 = np.array([15, 255, 255], dtype=np.uint8)
        lower_red_2 = np.array([160, 60, 50], dtype=np.uint8)
        upper_red_2 = np.array([179, 255, 255], dtype=np.uint8)
        mask_1 = cv2.inRange(hsv, lower_red_1, upper_red_1)
        mask_2 = cv2.inRange(hsv, lower_red_2, upper_red_2)
        red_mask = cv2.bitwise_or(mask_1, mask_2)

        if cv2.countNonZero(red_mask) == 0:
            return []

        red_mask = cv2.morphologyEx(red_mask, cv2.MORPH_CLOSE, np.ones((3, 3), np.uint8), iterations=1)
        red_mask = cv2.morphologyEx(red_mask, cv2.MORPH_OPEN, np.ones((2, 2), np.uint8), iterations=1)

        red_text = cv2.bitwise_and(work_image, work_image, mask=red_mask)
        red_gray = cv2.cvtColor(red_text, cv2.COLOR_BGR2GRAY)
        red_gray = cv2.normalize(red_gray, None, 0, 255, cv2.NORM_MINMAX)
        _, red_otsu = cv2.threshold(red_gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        variants: List[np.ndarray] = [red_gray, red_otsu, cv2.bitwise_not(red_otsu)]

        # Also try center crop on red region to focus on inner business name
        center_crop = OCRProcessor._crop_center(red_gray, ratio=0.65)
        if center_crop is not None and center_crop.size > 0:
            center_gray = OCRProcessor._enhance_image(center_crop)
            variants.append(center_gray)
            _, center_otsu = cv2.threshold(center_gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            variants.append(center_otsu)

        return variants

    @staticmethod
    def _build_image_variants(image: np.ndarray) -> List[np.ndarray]:
        """Build OCR-friendly variants to improve recall for seals and scanned images."""
        base = OCRProcessor._enhance_image(image)
        if base is None or base.size == 0:
            return [image]

        variants: List[np.ndarray] = [base]

        center_crop = OCRProcessor._crop_center(base, ratio=0.72)
        if center_crop is not None and center_crop.size > 0:
            variants.append(center_crop)

        _, otsu = cv2.threshold(base, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        variants.append(otsu)
        variants.append(cv2.bitwise_not(otsu))

        adaptive = cv2.adaptiveThreshold(
            base,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            11,
        )
        variants.append(adaptive)

        # Red stamp variants only (skip rotation variants for speed)
        variants.extend(OCRProcessor._build_red_stamp_variants(image))

        return variants

    @staticmethod
    def _normalize_text(text: str) -> str:
        folded = unicodedata.normalize("NFKD", (text or "").strip())
        folded = "".join(ch for ch in folded if not unicodedata.combining(ch))
        folded = folded.lower()
        return re.sub(r"[^a-z0-9]+", "", folded)

    @staticmethod
    def _restore_common_stamp_text(text: str) -> str:
        """Restore common Vietnamese stamp phrases that OCR often drops accents from."""
        if not text:
            return text

        replacements = [
            # Exact matches first
            (r"\bDOANH\s+NGHIEP\b", "DOANH NGHIỆP"),
            (r"\bDOANH\s+NGHIE\s+P\b", "DOANH NGHIỆP"),
            (r"\bDOANH\s+NGHIE\b", "DOANH NGHIỆP"),
            (r"\bTU\s+NHAN\b", "TƯ NHÂN"),
            (r"\bTUNHAN\b", "TƯ NHÂN"),
            (r"\bTU\s+NHAN\b", "TƯ NHÂN"),
            (r"\bTHIEN\s+Y\s+DAT\b", "THIÊN Ý ĐẠT"),
            (r"\bBIEN\s+HOA\b", "BIÊN HÒA"),
            (r"\bDONG\s+NAI\b", "ĐỒNG NAI"),
            (r"\bCONG\s+TY\b", "CÔNG TY"),
            (r"\bSAN\s+XUAT\b", "SẢN XUẤT"),
            (r"\bKINH\s+DOANH\b", "KINH DOANH"),
        ]

        normalized = text
        for pattern, replacement in replacements:
            normalized = re.sub(pattern, replacement, normalized, flags=re.IGNORECASE)
        return normalized

    @staticmethod
    def _collect_rapidocr_detections(
        reader: RapidOCR,
        variants: List[np.ndarray],
        min_confidence: float,
    ) -> List[Tuple[str, float, float, float]]:
        detections: List[Tuple[str, float, float, float]] = []
        for variant in variants:
            results, _elapsed = reader(variant)
            if not results:
                continue

            for detection in results:
                if not detection or len(detection) < 3:
                    continue

                box = detection[0] if len(detection) > 0 else None
                text = str(detection[1] or "").strip()
                try:
                    confidence = float(detection[2] or 0.0)
                except (TypeError, ValueError):
                    confidence = 0.0

                if not text or confidence < min_confidence:
                    continue

                x, y = 0.0, 0.0
                if isinstance(box, (list, tuple)) and box:
                    try:
                        x = float(sum(float(p[0]) for p in box) / len(box))
                        y = float(sum(float(p[1]) for p in box) / len(box))
                    except Exception:
                        x, y = 0.0, 0.0
                detections.append((text, confidence, y, x))
        return detections

    @staticmethod
    def _collect_easyocr_detections(
        easy_reader: Any,
        variants: List[np.ndarray],
        min_confidence: float,
    ) -> List[Tuple[str, float, float, float]]:
        detections: List[Tuple[str, float, float, float]] = []
        for variant in variants:
            try:
                results = easy_reader.readtext(variant, detail=1)
            except Exception:
                continue

            for detection in results or []:
                if not detection or len(detection) < 3:
                    continue

                box, text, confidence = detection[0], str(detection[1] or "").strip(), float(detection[2] or 0.0)
                if not text or confidence < min_confidence:
                    continue

                x, y = 0.0, 0.0
                if isinstance(box, (list, tuple)) and box:
                    try:
                        x = float(sum(float(p[0]) for p in box) / len(box))
                        y = float(sum(float(p[1]) for p in box) / len(box))
                    except Exception:
                        x, y = 0.0, 0.0

                detections.append((text, confidence, y, x))
        return detections

    @staticmethod
    def _is_circular_layout(positions: List[Tuple[float, float]]) -> bool:
        """Detect if text positions form a circular pattern."""
        if len(positions) < 3:
            return False
        
        # Calculate centroid
        cx = sum(x for x, y in positions) / len(positions)
        cy = sum(y for x, y in positions) / len(positions)
        
        # Calculate distances from center
        distances = [((x - cx) ** 2 + (y - cy) ** 2) ** 0.5 for x, y in positions]
        
        # Check if distances are relatively uniform (circular layout indicator)
        avg_dist = sum(distances) / len(distances)
        if avg_dist < 5:
            return False
        
        std_dev = (sum((d - avg_dist) ** 2 for d in distances) / len(distances)) ** 0.5
        variance_ratio = std_dev / avg_dist if avg_dist > 0 else 0
        
        # If std_dev is small relative to avg, it's likely circular (text radiates from center)
        return variance_ratio < 0.35

    @staticmethod
    def _sort_circular_text(items: List[Tuple[str, float, float, float]]) -> List[Tuple[str, float, float, float]]:
        """Sort circular text by angle from centroid (clockwise from top)."""
        if not items:
            return items
        
        # Calculate centroid
        cx = sum(x for _, _, _, x in items) / len(items)
        cy = sum(y for _, _, _, y in items) / len(items)
        
        def angle_from_center(item: Tuple[str, float, float, float]) -> float:
            _, _, y, x = item
            # atan2 gives angle; we adjust so top=0°, clockwise positive
            angle = np.arctan2(x - cx, cy - y)  # cy-y for top-to-bottom inversion
            return (angle + 2 * np.pi) % (2 * np.pi)  # Normalize to [0, 2π)
        
        return sorted(items, key=angle_from_center)

    @staticmethod
    def _finalize_detections(detections: List[Tuple[str, float, float, float]]) -> str:
        if not detections:
            return ""

        # First pass: dedup by normalized text, keep best confidence
        best_by_text: Dict[str, Tuple[str, float, float, float]] = {}
        for text, conf, y, x in detections:
            key = OCRProcessor._normalize_text(text)
            # Drop obvious OCR noise fragments.
            if not key or len(key) < 3:
                continue
            current = best_by_text.get(key)
            if current is None or conf > current[1]:
                best_by_text[key] = (text, conf, y, x)

        items_list = list(best_by_text.values())
        
        # Filter out obvious noise (single chars, very short lines that look like corruption)
        filtered_items: List[Tuple[str, float, float, float]] = []
        for text, conf, y, x in items_list:
            # Keep if: multi-word OR reasonable length >= 4 chars (after stripping)
            clean_text = text.strip()
            word_count = len(re.split(r"\s+", clean_text))
            if word_count >= 2 or len(clean_text) >= 4:
                filtered_items.append((text, conf, y, x))
        
        if not filtered_items:
            filtered_items = items_list  # Fallback: use all if nothing survives filter
        
        items_list = filtered_items
        
        # Sort by Y-coordinate (linear layout is most common)
        items_list.sort(key=lambda item: item[2])  # Sort by y
        
        rendered_lines: List[str] = []
        seen_normalized: set[str] = set()
        
        for text, confidence, _y, _x in items_list:
            clean_text = text.strip()
            normalized = OCRProcessor._normalize_text(clean_text)
            
            # Check for substring duplicates - keep the LONGER one
            skip = False
            to_remove: set[str] = set()
            
            for seen in list(seen_normalized):
                # If one is substring of the other, remove the shorter one
                if normalized in seen:
                    # Current is substring of seen, skip current
                    skip = True
                    break
                elif seen in normalized:
                    # Seen is substring of current, remove seen from output
                    to_remove.add(seen)
            
            if skip:
                continue
            
            # Remove shorter versions that are substrings
            if to_remove:
                # Re-filter rendered_lines to remove shorter versions
                rendered_lines = [
                    line for line in rendered_lines 
                    if OCRProcessor._normalize_text(line) not in to_remove
                ]
                seen_normalized = seen_normalized - to_remove
            
            seen_normalized.add(normalized)
            
            # Apply Vietnamese restoration
            restored = OCRProcessor._restore_common_stamp_text(clean_text)
            if restored:
                rendered_lines.append(restored)
        
        return "\n".join(rendered_lines)

    @staticmethod
    def extract_text_from_image(
        image: np.ndarray,
        config: OCRConfig = None
    ) -> str:
        """Extract text from image using RapidOCR with EasyOCR fallback when needed."""
        if config is None:
            config = OCRConfig()

        rapid_reader = OCRProcessor._get_reader(config)

        try:
            variants = OCRProcessor._build_image_variants(image)
            detections = OCRProcessor._collect_rapidocr_detections(
                rapid_reader,
                variants,
                min_confidence=config.confidence_threshold,
            )

            # For difficult stamps/images, fallback engine can recover missed characters.
            fallback_threshold = max(0.05, config.confidence_threshold - 0.1)
            if len(detections) < 3:
                easy_reader = OCRProcessor._get_easy_reader(config)
                if easy_reader is not None:
                    detections.extend(
                        OCRProcessor._collect_easyocr_detections(
                            easy_reader,
                            variants,
                            min_confidence=fallback_threshold,
                        )
                    )

            return OCRProcessor._restore_common_stamp_text(OCRProcessor._finalize_detections(detections))
        except Exception as e:
            print(f"[ERROR] Error in OCR: {e}")
            return ""

    @staticmethod
    def extract_images_from_pdf(
        pdf_path: str,
        dpi: int = 300
    ) -> List[Tuple[np.ndarray, int]]:
        """Extract rendered page images from PDF file."""
        result: List[Tuple[np.ndarray, int]] = []
        scale = max(dpi / 72.0, 1.0)

        def _collect(pdf_doc: Any) -> None:
            for page_num in range(len(pdf_doc)):
                page = pdf_doc.get_page(page_num)
                try:
                    bitmap = page.render(scale=scale)
                    img_array = bitmap.to_numpy()
                    if img_array is not None and img_array.size > 0:
                        result.append((img_array, page_num + 1))
                finally:
                    page.close()

        try:
            pdf = pdfium.PdfDocument(pdf_path)
            try:
                _collect(pdf)
            finally:
                pdf.close()
            return result
        except Exception as first_error:
            # Some Windows paths with Unicode may fail in native bindings; retry using bytes.
            try:
                pdf_bytes = Path(pdf_path).read_bytes()
                pdf = pdfium.PdfDocument(pdf_bytes)
                try:
                    _collect(pdf)
                finally:
                    pdf.close()
                return result
            except Exception as second_error:
                print(f"[ERROR] Error extracting images from PDF: {first_error}; fallback failed: {second_error}")
                return []
    
    @staticmethod
    def extract_images_from_docx(docx_path: str) -> List[Tuple[np.ndarray, str]]:
        """Extract images from DOCX file.
        
        Returns:
            List of tuples (image_array, image_name)
        """
        try:
            doc = DocxDocument(docx_path)
            images = []
            
            # Extract from document relationships
            for rel in doc.part.rels.values():
                target_part = getattr(rel, "target_part", None)
                if target_part is None:
                    continue

                content_type = getattr(target_part, "content_type", "") or ""
                if not content_type.startswith("image/"):
                    continue

                image_bytes = target_part.blob
                nparr = np.frombuffer(image_bytes, np.uint8)
                img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

                if img is not None:
                    image_name = Path(getattr(rel, "target_ref", "image")).name
                    images.append((img, image_name))
            
            return images
        except Exception as e:
            print(f"[ERROR] Error extracting images from DOCX: {e}")
            return []
    
    @staticmethod
    def process_pdf_with_ocr(
        pdf_path: str,
        config: OCRConfig = None,
        dpi: int = 300
    ) -> List[Document]:
        """Process PDF file and extract text from images using OCR."""
        if config is None:
            config = OCRConfig()

        documents = []

        try:
            image_list = OCRProcessor.extract_images_from_pdf(pdf_path, dpi=dpi)
            for img_array, page_num in image_list:
                text = OCRProcessor.extract_text_from_image(img_array, config)
                if text.strip():
                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "page": page_num,
                                "page_is_one_based": True,
                                "source_type": "pdf_ocr",
                                "extraction_method": "rapidocr+easyocr_fallback",
                            },
                        )
                    )

        except Exception as e:
            print(f"[ERROR] Error processing PDF with OCR: {e}")

        return documents
    
    @staticmethod
    def process_docx_with_ocr(
        docx_path: str,
        config: OCRConfig = None
    ) -> List[Document]:
        """Process DOCX file and extract text from embedded images using OCR."""
        if config is None:
            config = OCRConfig()

        documents = []

        try:
            image_list = OCRProcessor.extract_images_from_docx(docx_path)
            for img_array, image_name in image_list:
                text = OCRProcessor.extract_text_from_image(img_array, config)
                if text.strip():
                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "image_name": image_name,
                                "source_type": "docx_ocr",
                                "extraction_method": "rapidocr+easyocr_fallback",
                            },
                        )
                    )

        except Exception as e:
            print(f"[ERROR] Error processing DOCX with OCR: {e}")

        return documents

    @staticmethod
    def process_image_with_ocr(
        image_path: str,
        config: OCRConfig = None,
    ) -> List[Document]:
        """Process standalone image file and extract text using OCR."""
        if config is None:
            config = OCRConfig()

        try:
            image = cv2.imread(image_path, cv2.IMREAD_COLOR)
            if image is None:
                return []

            text = OCRProcessor.extract_text_from_image(image, config)
            if not text.strip():
                return []

            return [
                Document(
                    page_content=text,
                    metadata={
                        "image_name": Path(image_path).name,
                        "source_type": "image_ocr",
                        "extraction_method": "rapidocr+easyocr_fallback",
                    },
                )
            ]
        except Exception as e:
            print(f"[ERROR] Error processing image with OCR: {e}")
            return []


def extract_ocr_text_from_files(
    file_items: List[Tuple[str, bytes]],
    use_ocr: bool = False,
    ocr_config: OCRConfig = None
) -> List[Document]:
    """Extract text from PDF/DOCX/image files using OCR."""
    if not use_ocr:
        return []

    if ocr_config is None:
        ocr_config = OCRConfig()

    documents = []
    temp_paths: List[str] = []
    image_suffixes = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}

    try:
        for file_name, file_bytes in file_items:
            suffix = os.path.splitext(file_name)[1].lower()

            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
                temp_paths.append(tmp_path)

            if suffix == ".pdf":
                print(f"📄 Processing PDF with OCR: {file_name}")
                documents.extend(OCRProcessor.process_pdf_with_ocr(tmp_path, ocr_config))
            elif suffix == ".docx":
                print(f"📄 Processing DOCX with OCR: {file_name}")
                documents.extend(OCRProcessor.process_docx_with_ocr(tmp_path, ocr_config))
            elif suffix in image_suffixes:
                print(f"[*] Processing image with OCR: {file_name}")
                documents.extend(OCRProcessor.process_image_with_ocr(tmp_path, ocr_config))

    finally:
        for path in temp_paths:
            try:
                os.remove(path)
            except OSError:
                pass

    return documents
