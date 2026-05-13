"""RAG engine for SmartDoc AI using PDFPlumberLoader, FAISS, and Ollama."""

from __future__ import annotations

import os
import re
import tempfile
import time
import uuid
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from functools import lru_cache
from typing import Any, Dict, List, Optional, Sequence, Tuple

from langchain_classic.chains import RetrievalQA
from langchain_core.documents import Document
from langchain_core.callbacks import CallbackManagerForRetrieverRun
from langchain_core.prompts import PromptTemplate
from langchain_core.retrievers import BaseRetriever
from langchain_classic.retrievers import EnsembleRetriever
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.retrievers import BM25Retriever
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import OllamaLLM
from langdetect import LangDetectException, detect
from docx import Document as DocxDocument
from sentence_transformers import CrossEncoder

try:
    # python-bidi 0.4.x exposes get_display via bidi.algorithm, not bidi directly
    import bidi
    import bidi.algorithm
    if not hasattr(bidi, "get_display"):
        bidi.get_display = bidi.algorithm.get_display
    import easyocr as _easyocr
    _OCR_AVAILABLE = True
except Exception:
    _OCR_AVAILABLE = False

_OCR_TEXT_THRESHOLD = 100
_ocr_reader = None


def _get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        _ocr_reader = _easyocr.Reader(["vi"], gpu=False, verbose=False)
    return _ocr_reader


@dataclass
class RAGConfig:
    embedding_model: str = "sentence-transformers/paraphrase-multilingual-mpnet-base-v2"
    ollama_model: str = "qwen2.5:7b"
    chunk_size: int = 1000
    chunk_overlap: int = 100
    top_k: int = 3
    temperature: float = 0.1
    num_ctx: int = 4096
    num_predict: int = 512
    num_gpu: int = 0
    # Re-ranking configuration
    use_reranking: bool = False
    reranker_model: str = "cross-encoder/ms-marco-MiniLM-L-6-v2"
    rerank_top_n: int = 10
    use_hybrid_search: bool = False
    hybrid_semantic_weight: float = 0.7
    hybrid_keyword_weight: float = 0.3
    use_graph_rag: bool = False
    graph_expansion_hops: int = 1
    graph_max_related_chunks: int = 6
    use_self_rag: bool = False
    enable_query_rewriting: bool = False
    enable_multi_hop: bool = False
    max_hops: int = 3
    confidence_threshold: float = 0.7


def _validate_config(config: RAGConfig) -> None:
    if config.chunk_size <= 0:
        raise ValueError("chunk_size phải lớn hơn 0.")
    if config.chunk_overlap < 0:
        raise ValueError("chunk_overlap không được âm.")
    if config.chunk_overlap >= config.chunk_size:
        raise ValueError("chunk_overlap phải nhỏ hơn chunk_size.")
    if config.top_k <= 0:
        raise ValueError("top_k phải lớn hơn 0.")
    if config.graph_expansion_hops < 0:
        raise ValueError("graph_expansion_hops không được âm.")
    if config.graph_max_related_chunks < 0:
        raise ValueError("graph_max_related_chunks không được âm.")


def _default_upload_timestamp() -> str:
    return datetime.now(timezone.utc).isoformat()


def _merge_document_metadata(
    file_name: str,
    *,
    extra_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    ext = (os.path.splitext(file_name)[1] or "").lower().lstrip(".")
    merged = dict(extra_metadata or {})
    merged.setdefault("source", file_name)
    merged.setdefault("upload_time", _default_upload_timestamp())
    merged.setdefault("doc_id", str(uuid.uuid4()))
    merged.setdefault("file_type", ext or "unknown")
    return merged


def _normalize_filter_values(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, (list, tuple, set)):
        values = value
    else:
        values = [value]
    return [str(item).strip() for item in values if str(item).strip()]


def _sanitize_filter_metadata(
    filter_metadata: Optional[Dict[str, Any]] = None,
) -> Optional[Dict[str, Any]]:
    if not filter_metadata:
        return None

    sanitized: Dict[str, Any] = {}
    for key in ("source", "doc_id"):
        values = _normalize_filter_values(filter_metadata.get(key))
        if not values:
            continue
        sanitized[key] = values if len(values) > 1 else values[0]
    return sanitized or None


def metadata_matches_filter(
    metadata: Optional[Dict[str, Any]],
    filter_metadata: Optional[Dict[str, Any]] = None,
) -> bool:
    normalized_filter = _sanitize_filter_metadata(filter_metadata)
    if not normalized_filter:
        return True

    doc_metadata = metadata or {}
    for key, expected in normalized_filter.items():
        actual = doc_metadata.get(key)
        if isinstance(expected, list):
            if str(actual) not in {str(item) for item in expected}:
                return False
        elif str(actual) != str(expected):
            return False
    return True


def filter_documents_by_metadata(
    documents: Sequence[Document],
    filter_metadata: Optional[Dict[str, Any]] = None,
) -> List[Document]:
    normalized_filter = _sanitize_filter_metadata(filter_metadata)
    if not normalized_filter:
        return list(documents)
    return [
        doc
        for doc in documents
        if metadata_matches_filter(getattr(doc, "metadata", None), normalized_filter)
    ]


def _similarity_search_with_optional_filter(
    vectorstore: FAISS,
    query: str,
    *,
    k: int,
    filter_metadata: Optional[Dict[str, Any]] = None,
) -> List[Document]:
    search_filter = _sanitize_filter_metadata(filter_metadata)
    if search_filter:
        return vectorstore.similarity_search(query, k=k, filter=search_filter)
    return vectorstore.similarity_search(query, k=k)


def _ocr_image(pil_image: Any) -> str:
    if not _OCR_AVAILABLE:
        return ""
    try:
        import numpy as np
        results = _get_ocr_reader().readtext(np.array(pil_image.convert("RGB")), detail=0, paragraph=True)
        return "\n".join(results).strip()
    except Exception:
        return ""


def _render_pdf_page_to_image(tmp_path: str, page_idx: int, resolution: int = 200) -> Any:
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(tmp_path)
        try:
            bitmap = pdf[page_idx].render(scale=resolution / 72.0, rotation=0)
            return bitmap.to_pil()
        finally:
            pdf.close()
    except Exception:
        return None


def _words_to_text(word_list: list) -> str:
    lines: List[str] = []
    cur_line: List[str] = []
    cur_top: Optional[float] = None
    for w in sorted(word_list, key=lambda w: (round(w["top"] / 5) * 5, w["x0"])):
        if cur_top is None or abs(w["top"] - cur_top) > 5:
            if cur_line:
                lines.append(" ".join(cur_line))
            cur_line, cur_top = [w["text"]], w["top"]
        else:
            cur_line.append(w["text"])
    if cur_line:
        lines.append(" ".join(cur_line))
    return "\n".join(lines)


def _find_column_split(words: list, page_width: float) -> Optional[float]:
    """Find the gutter x-position between two columns by locating the lowest-density gap."""
    lo, hi = page_width * 0.30, page_width * 0.70
    bin_size = 5.0
    n_bins = max(1, int((hi - lo) / bin_size))
    density = [0] * n_bins
    for w in words:
        idx = int((w["x0"] - lo) / bin_size)
        if 0 <= idx < n_bins:
            density[idx] += 1
    avg = sum(density) / n_bins
    if avg == 0:
        return None
    min_val = min(density)
    if min_val > avg * 0.15:  # no clear gap → single column
        return None
    min_idx = density.index(min_val)
    return lo + (min_idx + 0.5) * bin_size


def _extract_page_text(page) -> str:
    words = page.extract_words(x_tolerance=2, y_tolerance=3)
    if not words:
        return (page.extract_text() or "").strip()

    split_x = _find_column_split(words, page.width)
    if split_x is not None:
        left = [w for w in words if w["x1"] <= split_x + 5]
        right = [w for w in words if w["x0"] >= split_x - 5]
        return (_words_to_text(left) + "\n\n" + _words_to_text(right)).strip()

    return _words_to_text(words).strip()


def _load_pdf_with_pdfplumber(tmp_path: str) -> List[Document]:
    import pdfplumber

    documents: List[Document] = []
    table_counter = 0
    with pdfplumber.open(tmp_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            text = _extract_page_text(page)

            if _OCR_AVAILABLE and len(text) < _OCR_TEXT_THRESHOLD:
                pil_img = _render_pdf_page_to_image(tmp_path, page_idx, resolution=300)
                if pil_img is not None:
                    ocr_text = _ocr_image(pil_img)
                    text = f"{text}\n{ocr_text}".strip() if text else ocr_text

            if text:
                documents.append(Document(page_content=text, metadata={"page": page_idx}))

            # Each table → separate document with sequential numbering
            for table in (page.extract_tables() or []):
                rows = [
                    " | ".join(str(cell).strip() if cell else "" for cell in row)
                    for row in table
                    if any(cell and str(cell).strip() for cell in row)
                ]
                if len(rows) < 2:  # skip chart labels / single-row artefacts
                    continue
                table_counter += 1
                table_text = (
                    f"Bảng {table_counter} (Trang {page_idx + 1}):\n"
                    + "\n".join(rows)
                )
                documents.append(Document(
                    page_content=table_text,
                    metadata={"page": page_idx, "table_index": table_counter},
                ))

    return documents


def load_documents_from_files(
    file_items: Sequence[Tuple[str, bytes]],
    document_metadata: Optional[Sequence[Optional[Dict[str, Any]]]] = None,
):
    """Load documents from supported document bytes (PDF, DOCX)."""

    documents = []
    temp_paths: List[str] = []
    try:
        metadata_items = list(document_metadata or [])
        for index, (file_name, file_bytes) in enumerate(file_items):
            base_metadata = metadata_items[index] if index < len(metadata_items) else None
            merged_metadata = _merge_document_metadata(file_name, extra_metadata=base_metadata)
            suffix = os.path.splitext(file_name)[1] or ".pdf"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            temp_paths.append(tmp_path)
            ext = suffix.lower()
            if ext == ".pdf":
                loaded_docs = _load_pdf_with_pdfplumber(tmp_path)
            elif ext == ".docx":
                loaded_docs = _load_docx_with_python_docx(tmp_path, file_name)
            else:
                raise ValueError(f"Định dạng file chưa được hỗ trợ: {file_name}")

            for doc in loaded_docs:
                doc.metadata = doc.metadata or {}
                doc.metadata.update(merged_metadata)

            documents.extend(loaded_docs)

        if not documents:
            raise ValueError("Không trích xuất được nội dung từ các file tải lên.")

        return documents
    finally:
        for path in temp_paths:
            try:
                os.remove(path)
            except OSError:
                pass


def _normalize_extracted_text(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines()]
    # Remove excessive blank lines while preserving paragraph boundaries.
    compact_lines: List[str] = []
    prev_blank = False
    for line in lines:
        if not line:
            if not prev_blank:
                compact_lines.append("")
            prev_blank = True
            continue

        compact_lines.append(line)
        prev_blank = False

    return "\n".join(compact_lines).strip()


def _load_docx_with_python_docx(path: str, file_name: str) -> List[Document]:
    import io
    from PIL import Image

    docx_file = DocxDocument(path)

    paragraphs = [p.text.strip() for p in docx_file.paragraphs if p.text.strip()]
    rows = [
        " | ".join(c.text.strip() for c in row.cells if c.text.strip())
        for table in docx_file.tables
        for row in table.rows
        if any(c.text.strip() for c in row.cells)
    ]
    images = []
    if _OCR_AVAILABLE:
        for rel in docx_file.part.rels.values():
            if "image" in rel.reltype:
                try:
                    text = _ocr_image(Image.open(io.BytesIO(rel.target_part.blob)))
                    if text:
                        images.append(text)
                except Exception:
                    pass

    normalized = _normalize_extracted_text("\n".join([*paragraphs, *rows, *images]))
    if not normalized:
        raise ValueError(f"Không trích xuất được nội dung từ file DOCX: {file_name}")

    return [Document(page_content=normalized, metadata={"page": 1, "page_is_one_based": True})]


def detect_main_language(documents: Sequence[Any]) -> str:
    collected: List[str] = []
    total_chars = 0
    for doc in documents:
        content = (getattr(doc, "page_content", "") or "").strip()
        if not content:
            continue

        remaining = 5000 - total_chars
        if remaining <= 0:
            break

        piece = content[:remaining]
        collected.append(piece)
        total_chars += len(piece)

    sample = "\n".join(collected).strip()
    if not sample:
        return "unknown"

    try:
        return detect(sample)
    except LangDetectException:
        return "unknown"


def detect_question_language(question: str) -> str:
    text = (question or "").strip()
    if not text:
        return "unknown"

    try:
        return detect(text)
    except LangDetectException:
        return "unknown"


def _normalize_lang_code(lang_code: str) -> str:
    code = (lang_code or "").strip().lower()
    if code.startswith("vi"):
        return "vi"
    if code.startswith("en"):
        return "en"
    return "unknown"


def resolve_response_language(question: str, doc_language: str = "unknown") -> str:
    question_lang = _normalize_lang_code(detect_question_language(question))
    document_lang = _normalize_lang_code(doc_language)

    # Very short questions are often ambiguous for language detection.
    if question_lang == "unknown" and len((question or "").split()) <= 5 and document_lang in {"vi", "en"}:
        return document_lang

    if question_lang in {"vi", "en"}:
        return question_lang
    if document_lang in {"vi", "en"}:
        return document_lang
    return "en"


def split_documents(documents, config: RAGConfig):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.chunk_size,
        chunk_overlap=config.chunk_overlap,
        add_start_index=True,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    for idx, chunk in enumerate(chunks, start=1):
        chunk.metadata = chunk.metadata or {}
        chunk.metadata["chunk_id"] = idx
        start_pos = chunk.metadata.get("start_index")
        if isinstance(start_pos, int):
            chunk.metadata["position_start"] = start_pos
            chunk.metadata["position_end"] = start_pos + len(chunk.page_content or "")
    return chunks


_GRAPH_STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "be", "by", "for", "from", "in", "into", "is", "it",
    "of", "on", "or", "that", "the", "this", "to", "was", "were", "with", "which", "what", "when",
    "where", "who", "why", "how", "do", "does", "did", "can", "could", "would", "should", "will",
    "you", "your", "we", "our", "they", "their", "them",
    "la", "là", "va", "và", "cua", "của", "cho", "trong", "ngoai", "ngoài", "nay", "này", "kia",
    "mot", "một", "nhung", "những", "cac", "các", "toi", "tôi", "ban", "bạn", "duoc", "được", "khong",
    "không", "co", "có", "se", "sẽ", "da", "đã", "ve", "về", "tai", "tại", "tu", "từ", "den", "đến",
}


def _graph_tokenize(text: str, max_terms: int = 24) -> set[str]:
    tokens = re.findall(r"[\w\-]{3,}", (text or "").lower(), flags=re.UNICODE)
    filtered = [tok for tok in tokens if tok not in _GRAPH_STOPWORDS and not tok.isdigit()]
    if not filtered:
        return set()

    counts: Dict[str, int] = {}
    for token in filtered:
        counts[token] = counts.get(token, 0) + 1

    ranked = sorted(counts.items(), key=lambda item: (-item[1], item[0]))
    return {token for token, _ in ranked[:max_terms]}


def _chunk_identifier(doc: Document, fallback_idx: int) -> int:
    metadata = doc.metadata or {}
    chunk_id = metadata.get("chunk_id")
    if isinstance(chunk_id, int):
        return chunk_id
    if isinstance(chunk_id, str) and chunk_id.isdigit():
        return int(chunk_id)
    return fallback_idx


def _build_chunk_graph(documents: Sequence[Document], min_overlap: int = 2) -> Dict[int, set[int]]:
    token_index: Dict[str, List[int]] = defaultdict(list)
    overlap_counts: Dict[int, Dict[int, int]] = defaultdict(dict)
    graph: Dict[int, set[int]] = defaultdict(set)

    doc_ids: List[int] = []
    for idx, doc in enumerate(documents, start=1):
        chunk_id = _chunk_identifier(doc, idx)
        doc_ids.append(chunk_id)
        tokens = _graph_tokenize(doc.page_content)
        for token in tokens:
            token_index[token].append(chunk_id)

    for chunk_id in doc_ids:
        graph.setdefault(chunk_id, set())

    for related_chunks in token_index.values():
        unique_chunks = sorted(set(related_chunks))
        if len(unique_chunks) < 2:
            continue

        for i in range(len(unique_chunks)):
            left = unique_chunks[i]
            for j in range(i + 1, len(unique_chunks)):
                right = unique_chunks[j]
                overlap_counts[left][right] = overlap_counts[left].get(right, 0) + 1
                overlap_counts[right][left] = overlap_counts[right].get(left, 0) + 1

    for source_chunk, neighbors in overlap_counts.items():
        for target_chunk, overlap in neighbors.items():
            if overlap >= min_overlap:
                graph[source_chunk].add(target_chunk)

    return graph


@lru_cache(maxsize=4)
def _get_embeddings(model_name: str) -> HuggingFaceEmbeddings:
    # Cache embedding model instance to avoid cold-loading on every rebuild.
    return HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


@lru_cache(maxsize=2)
def _get_cross_encoder(model_name: str) -> CrossEncoder:
    return CrossEncoder(model_name, max_length=512)


class RerankingRetriever(BaseRetriever):
    vectorstore: FAISS
    documents: List[Document]
    config: RAGConfig
    filter_metadata: Optional[Dict[str, Any]] = None

    class Config:
        arbitrary_types_allowed = True

    def _get_relevant_documents(
        self, query: str, *, run_manager: Optional[CallbackManagerForRetrieverRun] = None
    ) -> List[Document]:
        k = self.config.rerank_top_n if self.config.use_reranking else self.config.top_k
        search_filter = _sanitize_filter_metadata(self.filter_metadata)
        initial_docs = filter_documents_by_metadata(
            _similarity_search_with_optional_filter(self.vectorstore, query, k=k, filter_metadata=search_filter),
            search_filter,
        )
        if not self.config.use_reranking or not initial_docs:
            return initial_docs[: self.config.top_k]
        return rerank_documents(query, initial_docs, self.config)


class VectorSearchRetriever(BaseRetriever):
    vectorstore: FAISS
    documents: List[Document]
    config: RAGConfig
    filter_metadata: Optional[Dict[str, Any]] = None

    class Config:
        arbitrary_types_allowed = True

    def _get_relevant_documents(
        self, query: str, *, run_manager: Optional[CallbackManagerForRetrieverRun] = None
    ) -> List[Document]:
        search_filter = _sanitize_filter_metadata(self.filter_metadata)
        docs = _similarity_search_with_optional_filter(
            self.vectorstore,
            query,
            k=self.config.top_k,
            filter_metadata=search_filter,
        )
        return filter_documents_by_metadata(docs, search_filter)[: self.config.top_k]


class HybridSearchRetriever(BaseRetriever):
    vectorstore: FAISS
    documents: List[Document]
    config: RAGConfig
    filter_metadata: Optional[Dict[str, Any]] = None

    class Config:
        arbitrary_types_allowed = True

    def _get_relevant_documents(
        self, query: str, *, run_manager: Optional[CallbackManagerForRetrieverRun] = None
    ) -> List[Document]:
        filtered_documents = filter_documents_by_metadata(self.documents, self.filter_metadata)
        if not filtered_documents:
            return []

        candidate_k = self.config.rerank_top_n if self.config.use_reranking else self.config.top_k
        semantic_config = RAGConfig(**{**self.config.__dict__, "top_k": candidate_k})
        semantic_retriever = VectorSearchRetriever(
            vectorstore=self.vectorstore,
            documents=self.documents,
            config=semantic_config,
            filter_metadata=self.filter_metadata,
        )

        keyword_retriever = BM25Retriever.from_documents(filtered_documents)
        keyword_retriever.k = candidate_k

        hybrid_retriever = EnsembleRetriever(
            retrievers=[semantic_retriever, keyword_retriever],
            weights=[self.config.hybrid_semantic_weight, self.config.hybrid_keyword_weight],
        )

        combined_docs = list(hybrid_retriever.invoke(query) or [])
        combined_docs = filter_documents_by_metadata(combined_docs, self.filter_metadata)
        if not self.config.use_reranking or not combined_docs:
            return combined_docs[: self.config.top_k]

        return rerank_documents(query, combined_docs[:candidate_k], self.config)


class GraphRAGRetriever(BaseRetriever):
    vectorstore: FAISS
    documents: List[Document]
    config: RAGConfig
    filter_metadata: Optional[Dict[str, Any]] = None

    class Config:
        arbitrary_types_allowed = True

    def _get_relevant_documents(
        self, query: str, *, run_manager: Optional[CallbackManagerForRetrieverRun] = None
    ) -> List[Document]:
        search_filter = _sanitize_filter_metadata(self.filter_metadata)
        filtered_docs = filter_documents_by_metadata(self.documents, search_filter)
        if not filtered_docs:
            return []

        candidate_k = max(self.config.top_k, self.config.rerank_top_n)
        retrieval_config = RAGConfig(
            **{
                **self.config.__dict__,
                "top_k": candidate_k,
                "use_reranking": False,
            }
        )

        if self.config.use_hybrid_search:
            base_retriever: Any = HybridSearchRetriever(
                vectorstore=self.vectorstore,
                documents=self.documents,
                config=retrieval_config,
                filter_metadata=search_filter,
            )
        else:
            base_retriever = VectorSearchRetriever(
                vectorstore=self.vectorstore,
                documents=self.documents,
                config=retrieval_config,
                filter_metadata=search_filter,
            )

        initial_docs = list(base_retriever.invoke(query) or [])
        initial_docs = filter_documents_by_metadata(initial_docs, search_filter)
        if not initial_docs:
            return []

        doc_by_chunk_id = {
            _chunk_identifier(doc, idx): doc
            for idx, doc in enumerate(filtered_docs, start=1)
        }
        chunk_graph = _build_chunk_graph(filtered_docs)

        initial_chunk_ids = [
            _chunk_identifier(doc, idx)
            for idx, doc in enumerate(initial_docs, start=1)
        ]
        expanded_chunk_ids = set(initial_chunk_ids)
        frontier = set(initial_chunk_ids)

        for _ in range(self.config.graph_expansion_hops):
            next_frontier: set[int] = set()
            for chunk_id in frontier:
                neighbors = chunk_graph.get(chunk_id, set())
                next_frontier.update(neighbors)

            next_frontier -= expanded_chunk_ids
            if not next_frontier:
                break

            allowed = self.config.graph_max_related_chunks
            if allowed > 0:
                remaining = max(0, allowed - max(0, len(expanded_chunk_ids) - len(initial_chunk_ids)))
                if remaining == 0:
                    break
                ordered_neighbors = sorted(next_frontier)
                next_frontier = set(ordered_neighbors[:remaining])

            expanded_chunk_ids.update(next_frontier)
            frontier = next_frontier

        query_tokens = _graph_tokenize(query, max_terms=32)
        scored: List[Tuple[float, Document]] = []
        initial_rank = {
            chunk_id: rank
            for rank, chunk_id in enumerate(initial_chunk_ids)
        }

        for chunk_id in expanded_chunk_ids:
            doc = doc_by_chunk_id.get(chunk_id)
            if doc is None:
                continue

            # Combine lexical overlap and initial retrieval rank for stable ordering.
            overlap = len(query_tokens & _graph_tokenize(doc.page_content, max_terms=32))
            rank_bonus = max(0, candidate_k - initial_rank.get(chunk_id, candidate_k + 3))
            score = float(overlap * 2 + rank_bonus)
            scored.append((score, doc))

        scored.sort(key=lambda item: item[0], reverse=True)
        expanded_docs = [doc for _, doc in scored]
        if not self.config.use_reranking:
            return expanded_docs[: self.config.top_k]

        rerank_candidates = expanded_docs[:candidate_k]
        return rerank_documents(query, rerank_candidates, self.config)


def rerank_documents(
    query: str,
    documents: List[Document],
    config: RAGConfig,
    return_scores: bool = False,
) -> List[Document] | Tuple[List[Document], List[float]]:
    if not documents:
        return ([], []) if return_scores else []

    pairs = [[query, doc.page_content] for doc in documents]
    scores = _get_cross_encoder(config.reranker_model).predict(pairs)
    doc_score_pairs = sorted(zip(documents, scores), key=lambda x: x[1], reverse=True)

    top_k = min(config.top_k, len(documents))
    reranked_docs = [doc for doc, _ in doc_score_pairs[:top_k]]
    reranked_scores = [float(s) for _, s in doc_score_pairs[:top_k]]

    if return_scores:
        return reranked_docs, reranked_scores
    return reranked_docs


def build_vectorstore(chunks, config: RAGConfig) -> FAISS:
    if not chunks:
        raise ValueError("Không có chunks để tạo FAISS index.")

    embeddings = _get_embeddings(config.embedding_model)
    return FAISS.from_documents(chunks, embeddings)


def _build_prompt() -> PromptTemplate:
    template = """
Bạn là trợ lý AI chính xác và trung thực.
You are an accurate and honest AI assistant.

Yêu cầu:
1) Chỉ sử dụng thông tin trong phần Context để trả lời.
2) Nếu Context không đủ thông tin, chỉ được trả lời đúng một câu fallback đã nêu trong câu hỏi.
3) Trả lời theo đúng ngôn ngữ của câu hỏi người dùng (Việt hoặc Anh).
4) Trả lời ngắn gọn, đúng trọng tâm (3-4 câu).
5) Không tự suy diễn, không bịa thêm thông tin ngoài Context.

Context:
{context}

Question:
{question}

Answer:
""".strip()

    return PromptTemplate(template=template, input_variables=["context", "question"])


def build_qa_chain(vectorstore: FAISS, chunks: List[Document], config: RAGConfig) -> RetrievalQA:
    llm = OllamaLLM(
        model=config.ollama_model,
        temperature=config.temperature,
        num_ctx=config.num_ctx,
        num_predict=config.num_predict,
        num_gpu=config.num_gpu,
        keep_alive="30m",
    )
    prompt = _build_prompt()

    if config.use_graph_rag:
        retriever = GraphRAGRetriever(vectorstore=vectorstore, documents=chunks, config=config)
    elif config.use_hybrid_search:
        retriever = HybridSearchRetriever(vectorstore=vectorstore, documents=chunks, config=config)
    elif config.use_reranking:
        retriever = RerankingRetriever(vectorstore=vectorstore, documents=chunks, config=config)
    else:
        retriever = VectorSearchRetriever(vectorstore=vectorstore, documents=chunks, config=config)

    return RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=retriever,
        return_source_documents=True,
        chain_type_kwargs={"prompt": prompt},
    )


def build_rag_pipeline(
    file_items: Sequence[Tuple[str, bytes]],
    config: RAGConfig,
    document_metadata: Optional[Sequence[Optional[Dict[str, Any]]]] = None,
) -> Tuple[RetrievalQA, str, int]:
    _validate_config(config)
    documents = load_documents_from_files(file_items, document_metadata=document_metadata)
    language = detect_main_language(documents)
    chunks = split_documents(documents, config)

    vectorstore = build_vectorstore(chunks, config)
    qa_chain = build_qa_chain(vectorstore, chunks, config)
    return qa_chain, language, len(chunks)


def _format_recent_history(conversation_history: Optional[Sequence[Tuple[str, str]]]) -> str:
    if not conversation_history:
        return "(none)"

    recent = list(conversation_history)[-6:]
    lines: List[str] = []
    for role, content in recent:
        label = "User" if role == "user" else "Assistant"
        lines.append(f"{label}: {content}")
    return "\n".join(lines)


def _looks_like_followup(question: str) -> bool:
    text = (question or "").strip().lower()
    if not text:
        return False

    followup_markers = [
        "nó",
        "cái đó",
        "điều đó",
        "ý trên",
        "tiếp theo",
        "them",
        "thêm",
        "what about",
        "how about",
        "that one",
        "those",
        "more details",
        "explain more",
        "continue",
    ]
    if any(marker in text for marker in followup_markers):
        return True

    # Short, ambiguous queries are often follow-up questions.
    return len(text.split()) <= 8


def contextualize_question(
    question: str,
    llm: OllamaLLM,
    conversation_history: Optional[Sequence[Tuple[str, str]]] = None,
) -> Dict[str, Any]:
    if not conversation_history:
        return {
            "original": question,
            "standalone": question,
            "used_contextualization": False,
            "followup_detected": False,
        }

    followup_detected = _looks_like_followup(question)
    recent_history = list(conversation_history)[-8:]
    history_text = _format_recent_history(recent_history)

    prompt = f"""
You rewrite conversational follow-up questions for RAG retrieval.
Bạn viết lại câu hỏi follow-up cho truy hồi RAG.

Recent conversation:
{history_text}

New user question:
{question}

Requirements:
1) If the new question is a follow-up, rewrite it into ONE standalone query with enough context.
2) If it is already standalone, keep its core meaning unchanged.
3) Keep the same language as the user's question (Vietnamese or English).
4) Return ONLY the rewritten standalone query, no explanation.

Standalone query:
""".strip()

    try:
        rewritten = (llm.invoke(prompt) or "").strip()
        if not rewritten:
            rewritten = question
        return {
            "original": question,
            "standalone": rewritten,
            "used_contextualization": True,
            "followup_detected": followup_detected,
        }
    except Exception:
        return {
            "original": question,
            "standalone": question,
            "used_contextualization": False,
            "followup_detected": followup_detected,
        }


def _deduplicate_documents(documents: List[Document]) -> List[Document]:
    seen: set[Tuple[str, str, Any]] = set()
    unique_docs: List[Document] = []
    for doc in documents:
        meta = doc.metadata or {}
        key = (
            str(meta.get("source", "")),
            str(meta.get("chunk_id", "")),
            meta.get("start_index", meta.get("position_start")),
        )
        if key in seen:
            continue
        seen.add(key)
        unique_docs.append(doc)
    return unique_docs


def _build_context_from_documents(documents: List[Document], max_chars: int = 3600) -> str:
    parts: List[str] = []
    current_len = 0
    for idx, doc in enumerate(documents, start=1):
        meta = doc.metadata or {}
        source = meta.get("source", "unknown")
        page = _normalize_page_number(meta.get("page", meta.get("page_number", "?")), meta)
        raw_text = (doc.page_content or "").strip()
        if len(raw_text) > 1200:
            raw_text = raw_text[:1200].rstrip() + "..."

        section = f"[Doc {idx} | Source: {source} | Page: {page}]\n{raw_text}"
        if not section.strip():
            continue

        projected = current_len + len(section) + 2
        if projected > max_chars:
            remaining = max_chars - current_len
            if remaining > 80:
                parts.append(section[:remaining].rstrip())
            break

        parts.append(section)
        current_len = projected

    return "\n\n".join(parts).strip()


def _looks_truncated_answer(text: str) -> bool:
    cleaned = (text or "").strip()
    if not cleaned:
        return False

    if cleaned.endswith((".", "!", "?", "\"", "'", "”", "’", "))", "]")):
        return False

    lower_tail = cleaned[-40:].lower()
    tail_markers = [
        " and",
        " or",
        " because",
        " such as",
        " bao gồm",
        " và",
        " hoặc",
        " vì",
        " như",
    ]
    if any(lower_tail.endswith(marker) for marker in tail_markers):
        return True
    return len(cleaned) >= 120


def _complete_truncated_answer(
    llm: OllamaLLM,
    partial_answer: str,
    question: str,
    context_text: str,
    response_language: str,
) -> str:
    completion_prompt = f"""
Continue the unfinished answer below and complete it naturally.

Language: {response_language}
Question: {question}
Context:
{context_text[:1800]}

Current partial answer:
{partial_answer}

Requirements:
1) Continue only from where it stopped; do not restart from the beginning.
2) Use only facts from Context.
3) Add at most 2 short sentences and end with a complete sentence.
4) Return only the continuation text.

Continuation:
""".strip()

    try:
        continuation = (llm.invoke(completion_prompt) or "").strip()
    except Exception:
        continuation = ""

    if not continuation:
        return partial_answer

    joined = f"{partial_answer.rstrip()} {continuation.lstrip()}".strip()
    return joined


def _maybe_translate_answer_to_target_language(
    llm: OllamaLLM,
    answer: str,
    target_lang: str,
) -> str:
    text = (answer or "").strip()
    target = _normalize_lang_code(target_lang)
    if not text or target not in {"vi", "en"}:
        return text

    try:
        detected = _normalize_lang_code(detect(text))
    except Exception:
        detected = "unknown"

    if detected == target:
        return text

    if target == "vi":
        prompt = f"""
Dich doan tra loi sau sang tieng Viet tu nhien.
Giu nguyen y nghia, ten rieng, so lieu va cac thong tin ky thuat.
Chi tra ve ban dich, khong them giai thich.

Noi dung:
{text}

Ban dich tieng Viet:
""".strip()
    else:
        prompt = f"""
Translate the following answer into natural English.
Preserve all facts, names, numbers, and technical details.
Return only the translated answer, no explanation.

Content:
{text}

English translation:
""".strip()

    try:
        translated = (llm.invoke(prompt) or "").strip()
        return translated or text
    except Exception:
        return text


def _is_uncertain_non_answer(text: str) -> bool:
    cleaned = (text or "").strip().casefold()
    if not cleaned:
        return True

    patterns = [
        "không thể hiểu",
        "khong the hieu",
        "cung cấp thêm thông tin",
        "cung cap them thong tin",
        "mô tả chi tiết",
        "mo ta chi tiet",
        "bạn có thể cung cấp",
        "ban co the cung cap",
        "i cannot understand",
        "cannot understand",
        "please provide more information",
        "can you provide more",
    ]
    return any(p in cleaned for p in patterns)


def _regenerate_with_strict_answering(
    llm: OllamaLLM,
    question: str,
    context_text: str,
    response_language: str,
) -> str:
    prompt = f"""
You must answer using only the provided excerpts.

Language: {response_language}
Question: {question}

Excerpts:
{context_text}

Rules:
1) Do not ask the user for more information.
2) Do not say you cannot understand.
3) Provide the best possible concise summary from the excerpts.
4) If evidence is limited, state that briefly but still answer from available excerpts.
5) 3-5 sentences.

Answer:
""".strip()

    try:
        return (llm.invoke(prompt) or "").strip()
    except Exception:
        return ""


def _generate_answer_from_documents(
    llm: OllamaLLM,
    question: str,
    documents: List[Document],
    response_lang: str,
    doc_language: str = "unknown",
    conversation_history: Optional[Sequence[Tuple[str, str]]] = None,
) -> str:
    lang_code = _normalize_lang_code(response_lang)
    response_language = "Vietnamese" if lang_code == "vi" else "English"
    unknown_phrase = _unknown_phrase_for_language(lang_code)
    chat_history = _format_recent_history(conversation_history)
    context_text = _build_context_from_documents(documents)

    if lang_code == "vi":
        prompt = f"""
Bạn là trợ lý AI chính xác và trung thực.

Lịch sử hội thoại gần đây:
{chat_history}

Ngôn ngữ tài liệu tham khảo: {doc_language}

Context:
{context_text}

Câu hỏi hiện tại:
{question}

Yêu cầu bắt buộc:
1) Chỉ dùng thông tin có trong Context.
2) Có thể dùng lịch sử hội thoại để hiểu tham chiếu như "nó", "ý trước".
3) Nếu Context không đủ thông tin, chỉ trả về đúng một câu: {unknown_phrase}
4) Chỉ trả lời bằng tiếng Việt.
5) Trả lời ngắn gọn, đầy đủ ý chính.

Trả lời:
""".strip()
    else:
        prompt = f"""
You are an accurate and honest AI assistant.

Recent conversation:
{chat_history}

Document language hint: {doc_language}

Context:
{context_text}

Current user question:
{question}

Mandatory rules:
1) Use only facts from Context.
2) You may use conversation history only to resolve references.
3) If Context is insufficient, return exactly one sentence: {unknown_phrase}
4) Respond only in English.
5) Keep the answer concise but complete.

Answer:
""".strip()

    answer = (llm.invoke(prompt) or "").strip()

    if context_text and _is_uncertain_non_answer(answer):
        retried = _regenerate_with_strict_answering(
            llm,
            question,
            context_text,
            response_language,
        )
        if retried:
            answer = retried

    if _looks_truncated_answer(answer):
        answer = _complete_truncated_answer(
            llm,
            answer,
            question,
            context_text,
            response_language,
        )

    answer = _maybe_translate_answer_to_target_language(llm, answer, lang_code)
    return _normalize_unknown_answer(answer, lang_code)


def _unknown_phrase_for_language(lang_code: str) -> str:
    return "Tôi không biết" if lang_code == "vi" else "I don't know"


def _normalize_unknown_answer(answer: str, lang_code: str) -> str:
    fallback = _unknown_phrase_for_language(lang_code)
    cleaned = (answer or "").strip()
    if not cleaned:
        return fallback

    normalized = cleaned.casefold().strip()
    fallback_norm = fallback.casefold().strip()
    if normalized.startswith(fallback_norm):
        tail = normalized[len(fallback_norm):].strip(" .,:;!?-\n\t")
        if tail:
            return fallback
        return fallback

    lowered = cleaned.casefold()
    unknown_patterns = [
        "khong du thong tin",
        "không đủ thông tin",
        "tai lieu khong co",
        "tài liệu không có",
        "không thể hiểu",
        "khong the hieu",
        "cung cấp thêm thông tin",
        "cung cap them thong tin",
        "bạn có thể cung cấp",
        "ban co the cung cap",
        "not enough information",
        "does not contain",
        "cannot find in context",
        "insufficient context",
        "cannot understand",
        "please provide more information",
        "i don't know",
        "toi khong biet",
        "tôi không biết",
        "我不知道",
    ]

    if any(pattern in lowered for pattern in unknown_patterns):
        return fallback

    return cleaned


def _normalize_page_number(page_value: Any, metadata: Dict[str, Any]) -> Any:
    if isinstance(page_value, int):
        if metadata.get("page_is_one_based"):
            return page_value
        return page_value + 1 if page_value >= 0 else page_value
    return page_value


def _build_highlight_text(text: str, max_chars: int = 260) -> str:
    cleaned = (text or "").strip()
    if len(cleaned) <= max_chars:
        return cleaned
    return cleaned[:max_chars].rstrip() + "..."


def _extract_llm_from_qa_chain(qa_chain: RetrievalQA) -> OllamaLLM:
    combine_chain = getattr(qa_chain, "combine_documents_chain", None)
    if combine_chain is not None:
        llm_chain = getattr(combine_chain, "llm_chain", None)
        llm = getattr(llm_chain, "llm", None) if llm_chain is not None else None
        if llm is not None:
            return llm
        llm = getattr(combine_chain, "llm", None)
        if llm is not None:
            return llm

    qa_llm_chain = getattr(qa_chain, "llm_chain", None)
    llm = getattr(qa_llm_chain, "llm", None) if qa_llm_chain is not None else None
    if llm is not None:
        return llm

    llm = getattr(qa_chain, "llm", None)
    if llm is not None:
        return llm

    raise AttributeError("Unable to resolve LLM from RetrievalQA.")


def _retrieve_documents(retriever: Any, query: str) -> List[Document]:
    if hasattr(retriever, "invoke"):
        return list(retriever.invoke(query) or [])
    if hasattr(retriever, "get_relevant_documents"):
        return retriever.get_relevant_documents(query)
    raise AttributeError("Retriever does not support `invoke` or `get_relevant_documents`.")


def _set_retriever_filter(
    retriever: Any,
    filter_metadata: Optional[Dict[str, Any]] = None,
) -> Optional[Dict[str, Any]]:
    if not hasattr(retriever, "filter_metadata"):
        return None
    previous_filter = getattr(retriever, "filter_metadata", None)
    retriever.filter_metadata = _sanitize_filter_metadata(filter_metadata)
    return previous_filter


def _restore_retriever_filter(retriever: Any, previous_filter: Optional[Dict[str, Any]]) -> None:
    if hasattr(retriever, "filter_metadata"):
        retriever.filter_metadata = previous_filter


def _translate_query_for_retrieval(
    query: str,
    source_lang: str,
    target_lang: str,
    llm: OllamaLLM,
) -> str:
    if not query.strip() or source_lang == target_lang:
        return query

    if source_lang == "vi" and target_lang == "en":
        prompt = f"""
Translate this Vietnamese search query into natural English for semantic document retrieval.
Keep meaning, named entities, numbers, and constraints unchanged.
Return only the translated query.

Vietnamese query:
{query}

English query:
""".strip()
    elif source_lang == "en" and target_lang == "vi":
        prompt = f"""
Dich cau truy van tim kiem tieng Anh sau sang tieng Viet de truy hoi tai lieu.
Giu nguyen y nghia, ten rieng, so lieu va rang buoc.
Chi tra ve cau dich, khong giai thich.

English query:
{query}

Vietnamese query:
""".strip()
    else:
        return query

    try:
        translated = (llm.invoke(prompt) or "").strip()
        return translated or query
    except Exception:
        return query


def _build_retrieval_queries(
    base_query: str,
    question_lang: str,
    doc_language: str,
    llm: OllamaLLM,
) -> List[str]:
    queries = [base_query]
    q_lang = _normalize_lang_code(question_lang)
    d_lang = _normalize_lang_code(doc_language)

    if q_lang in {"vi", "en"} and d_lang in {"vi", "en"} and q_lang != d_lang:
        translated = _translate_query_for_retrieval(base_query, q_lang, d_lang, llm)
        if translated and translated.strip() and translated.strip() != base_query.strip():
            queries.append(translated.strip())

    unique: List[str] = []
    seen: set[str] = set()
    for q in queries:
        key = q.strip().casefold()
        if key and key not in seen:
            seen.add(key)
            unique.append(q.strip())
    return unique


def _retrieve_documents_cross_language(
    retriever: Any,
    base_query: str,
    question_lang: str,
    doc_language: str,
    llm: OllamaLLM,
) -> Tuple[List[Document], List[str]]:
    queries = _build_retrieval_queries(base_query, question_lang, doc_language, llm)

    merged_docs: List[Document] = []
    for q in queries:
        merged_docs.extend(_retrieve_documents(retriever, q))

    merged_docs = _deduplicate_documents(list(merged_docs or []))

    # If retriever includes config (RerankingRetriever), re-rank merged candidates.
    cfg = getattr(retriever, "config", None)
    if cfg is not None and getattr(cfg, "use_reranking", False) and merged_docs:
        rerank_query = queries[0]
        merged_docs = rerank_documents(rerank_query, merged_docs, cfg)

    return merged_docs, queries


def rewrite_query(
    original_query: str,
    llm: OllamaLLM,
    language: str = "unknown",
) -> Dict[str, Any]:
    prompt_template = """
You are a query optimizer for RAG retrieval.
Bạn là công cụ tối ưu câu hỏi cho truy hồi RAG.

Original query:
{query}

Requirements:
1) Preserve the original meaning.
2) Keep the same language as the original query.
3) Make entities, constraints, and intent explicit for better retrieval.
4) Return ONLY one rewritten query, no explanation.

Rewritten query:
""".strip()

    try:
        prompt = prompt_template.format(query=original_query)
        rewritten = llm.invoke(prompt).strip()

        return {
            "original": original_query,
            "rewritten": rewritten,
            "used_rewriting": True,
            "language": language,
        }
    except Exception:
        return {
            "original": original_query,
            "rewritten": original_query,
            "used_rewriting": False,
            "language": language,
        }


def evaluate_answer_quality(
    question: str,
    answer: str,
    context_documents: List[Document],
    llm: OllamaLLM,
) -> Dict[str, Any]:
    context_text = "\n\n".join([
        f"Document {i+1}: {doc.page_content[:200]}..."
        for i, doc in enumerate(context_documents[:3])
    ])

    evaluation_prompt = """
You are an answer quality evaluator.
Bạn là người đánh giá chất lượng câu trả lời.

Question / Câu hỏi: {question}

Context:
{context}

Answer to evaluate:
{answer}

Score from 0-10 for each criterion:
1. Grounding (0-10): Is the answer supported by context?
2. Relevance (0-10): Does it answer the question?
3. Completeness (0-10): Is it sufficiently complete?

Return only 3 numbers, one per line. Example:
8
9
7
""".strip()

    try:
        prompt = evaluation_prompt.format(
            question=question,
            context=context_text,
            answer=answer
        )

        evaluation_text = llm.invoke(prompt).strip()

        lines = [line.strip() for line in evaluation_text.split('\n') if line.strip()]
        scores = []
        for line in lines[:3]:
            try:
                scores.append(min(10.0, max(0.0, float(line))))
            except ValueError:
                numbers = re.findall(r'\d+\.?\d*', line)
                if numbers:
                    scores.append(min(10.0, max(0.0, float(numbers[0]))))

        while len(scores) < 3:
            scores.append(5.0)

        grounding, relevance, completeness = scores[:3]
        confidence = (grounding + relevance + completeness) / 30.0

        return {
            "grounding_score": grounding,
            "relevance_score": relevance,
            "completeness_score": completeness,
            "confidence": confidence,
            "passed": confidence >= 0.7,
        }

    except Exception as e:
        return {
            "grounding_score": 7.0,
            "relevance_score": 7.0,
            "completeness_score": 7.0,
            "confidence": 0.7,
            "passed": True,
            "error": str(e),
        }


def multi_hop_retrieval(
    original_query: str,
    vectorstore: FAISS,
    llm: OllamaLLM,
    config: RAGConfig,
    max_hops: int = 3,
    filter_metadata: Optional[Dict[str, Any]] = None,
    retriever: Optional[Any] = None,
) -> Dict[str, Any]:
    all_documents = []
    hop_history = []
    current_query = original_query
    search_filter = _sanitize_filter_metadata(filter_metadata)

    for hop in range(max_hops):
        # Retrieve documents for current query
        if retriever is not None:
            previous_filter = _set_retriever_filter(retriever, search_filter)
            try:
                hop_docs = _retrieve_documents(retriever, current_query)
            finally:
                _restore_retriever_filter(retriever, previous_filter)
        elif config.use_reranking:
            candidates = _similarity_search_with_optional_filter(
                vectorstore,
                current_query,
                k=config.rerank_top_n,
                filter_metadata=search_filter,
            )
            candidates = filter_documents_by_metadata(candidates, search_filter)
            hop_docs = rerank_documents(current_query, candidates, config)
        else:
            hop_docs = _similarity_search_with_optional_filter(
                vectorstore,
                current_query,
                k=config.top_k,
                filter_metadata=search_filter,
            )
            hop_docs = filter_documents_by_metadata(hop_docs, search_filter)

        all_documents.extend(hop_docs)

        context_text = "\n\n".join([doc.page_content for doc in all_documents])

        answer_prompt = f"""
    Answer the question using context below.
    Trả lời câu hỏi dựa trên context bên dưới.

    Original question: {original_query}
    Current question: {current_query}

Context:
{context_text[:2000]}

    If context is enough, provide the answer.
    If not enough, clearly state what information is still needed.

    Answer:
""".strip()

        partial_answer = llm.invoke(answer_prompt).strip()

        need_more_info = any(phrase in partial_answer.lower() for phrase in [
            "cần thêm", "thiếu thông tin", "không đủ", "need more", "insufficient"
        ])

        hop_history.append({
            "hop": hop + 1,
            "query": current_query,
            "num_docs": len(hop_docs),
            "partial_answer": partial_answer,
            "need_more_info": need_more_info,
        })

        if not need_more_info:
            return {
                "final_answer": partial_answer,
                "total_hops": hop + 1,
                "all_documents": all_documents,
                "hop_history": hop_history,
                "completed": True,
            }

        if hop < max_hops - 1:
            followup_prompt = f"""
Original question: {original_query}
Current findings: {partial_answer}

Create ONE follow-up retrieval query to fetch missing facts.
Return only the query text, no explanation.

Follow-up query:
""".strip()

            current_query = llm.invoke(followup_prompt).strip()

    return {
        "final_answer": partial_answer,
        "total_hops": max_hops,
        "all_documents": all_documents,
        "hop_history": hop_history,
        "completed": False,
    }


def compare_retrieval_methods(
    vectorstore: FAISS,
    documents: Sequence[Document],
    query: str,
    config: RAGConfig,
    filter_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    start_time = time.time()
    search_filter = _sanitize_filter_metadata(filter_metadata)

    vector_retriever = VectorSearchRetriever(
        vectorstore=vectorstore, documents=list(documents), config=config, filter_metadata=search_filter,
    )
    hybrid_retriever = HybridSearchRetriever(
        vectorstore=vectorstore,
        documents=list(documents),
        config=RAGConfig(**{**config.__dict__, "use_hybrid_search": True, "use_reranking": False}),
        filter_metadata=search_filter,
    )

    bi_encoder_start = time.time()
    bi_encoder_docs = vector_retriever.invoke(query)
    bi_encoder_time = time.time() - bi_encoder_start

    hybrid_start = time.time()
    hybrid_docs = hybrid_retriever.invoke(query)
    hybrid_time = time.time() - hybrid_start

    retrieval_start = time.time()
    candidate_docs = filter_documents_by_metadata(
        _similarity_search_with_optional_filter(vectorstore, query, k=config.rerank_top_n, filter_metadata=search_filter),
        search_filter,
    )
    retrieval_time = time.time() - retrieval_start

    rerank_start = time.time()
    reranked_docs, rerank_scores = rerank_documents(query, candidate_docs, config, return_scores=True)
    rerank_time = time.time() - rerank_start

    total_time = time.time() - start_time

    return {
        "bi_encoder": {
            "docs": bi_encoder_docs,
            "time_ms": round(bi_encoder_time * 1000, 2),
            "num_docs": len(bi_encoder_docs),
        },
        "hybrid": {
            "docs": hybrid_docs,
            "time_ms": round(hybrid_time * 1000, 2),
            "num_docs": len(hybrid_docs),
            "weights": [config.hybrid_semantic_weight, config.hybrid_keyword_weight],
        },
        "cross_encoder": {
            "docs": reranked_docs,
            "scores": rerank_scores,
            "time_ms": round((retrieval_time + rerank_time) * 1000, 2),
            "retrieval_time_ms": round(retrieval_time * 1000, 2),
            "rerank_time_ms": round(rerank_time * 1000, 2),
            "num_candidates": len(candidate_docs),
            "num_final": len(reranked_docs),
        },
        "comparison": {
            "total_time_ms": round(total_time * 1000, 2),
            "speedup": round(bi_encoder_time / (retrieval_time + rerank_time), 2),
            "overhead_ms": round((retrieval_time + rerank_time - bi_encoder_time) * 1000, 2),
            "hybrid_overhead_ms": round((hybrid_time - bi_encoder_time) * 1000, 2),
            "hybrid_vs_vector_ratio": round(hybrid_time / bi_encoder_time, 2) if bi_encoder_time else None,
        },
    }


def _format_source_documents(documents: Sequence[Document], top_k: int) -> List[Dict[str, Any]]:
    sources: List[Dict[str, Any]] = []
    for doc in list(documents)[:top_k]:
        metadata = doc.metadata or {}
        context_text = (doc.page_content or "").strip()
        start_pos = metadata.get("position_start", metadata.get("start_index"))
        end_pos = metadata.get("position_end")
        if end_pos is None and isinstance(start_pos, int):
            end_pos = start_pos + len(context_text)

        page_value = metadata.get("page", metadata.get("page_number", "?"))
        page_value = _normalize_page_number(page_value, metadata)
        highlight_text = _build_highlight_text(context_text)
        sources.append(
            {
                "chunk_id": metadata.get("chunk_id", "?"),
                "source": metadata.get("source", "unknown"),
                "doc_id": metadata.get("doc_id"),
                "upload_time": metadata.get("upload_time"),
                "file_type": metadata.get("file_type"),
                "page": page_value,
                "position_start": start_pos,
                "position_end": end_pos,
                "preview": highlight_text.replace("\n", " "),
                "highlight_text": highlight_text,
                "context_text": context_text,
            }
        )
    return sources


def ask_question_with_self_rag(
    qa_chain: RetrievalQA,
    question: str,
    vectorstore: FAISS,
    config: RAGConfig,
    conversation_history: Optional[Sequence[Tuple[str, str]]] = None,
    filter_metadata: Optional[Dict[str, Any]] = None,
    doc_language: str = "unknown",
) -> Tuple[str, List[Dict[str, Any]], Dict[str, Any]]:
    metadata = {
        "used_self_rag": config.use_self_rag,
        "query_rewriting": config.enable_query_rewriting,
        "multi_hop": config.enable_multi_hop,
        "filter_metadata": _sanitize_filter_metadata(filter_metadata),
    }

    llm = _extract_llm_from_qa_chain(qa_chain)

    contextualized = contextualize_question(question, llm, conversation_history)
    working_query = contextualized["standalone"]
    metadata["contextualization"] = contextualized

    if config.use_self_rag and config.enable_query_rewriting:
        rewrite_result = rewrite_query(
            working_query,
            llm,
            language=resolve_response_language(question, doc_language),
        )
        working_query = rewrite_result["rewritten"]
        metadata["query_rewrite"] = rewrite_result
    else:
        metadata["query_rewrite"] = None

    question_lang = resolve_response_language(question, doc_language)

    if config.use_self_rag and config.enable_multi_hop:
        multi_hop_result = multi_hop_retrieval(
            working_query,
            vectorstore,
            llm,
            config,
            max_hops=config.max_hops,
            filter_metadata=filter_metadata,
            retriever=qa_chain.retriever,
        )
        retrieved_docs = _deduplicate_documents(multi_hop_result["all_documents"])
        metadata["multi_hop"] = multi_hop_result
        metadata["retrieval_queries"] = [working_query]
    else:
        retriever = qa_chain.retriever
        previous_filter = _set_retriever_filter(retriever, filter_metadata)
        try:
            retrieved_docs, used_queries = _retrieve_documents_cross_language(
                retriever,
                working_query,
                question_lang,
                doc_language,
                llm,
            )
        finally:
            _restore_retriever_filter(retriever, previous_filter)

        metadata["multi_hop"] = None
        metadata["retrieval_queries"] = used_queries

    answer = _generate_answer_from_documents(
        llm,
        question,
        list(retrieved_docs or []),
        question_lang,
        doc_language,
        conversation_history,
    )

    if config.use_self_rag:
        evaluation = evaluate_answer_quality(
            question,
            answer,
            list(retrieved_docs or []),
            llm,
        )
        metadata["evaluation"] = evaluation
        if evaluation["confidence"] < config.confidence_threshold:
            metadata["low_confidence_warning"] = True
    else:
        metadata["evaluation"] = None

    return answer, _format_source_documents(list(retrieved_docs or []), config.top_k), metadata


def ask_question(
    qa_chain: RetrievalQA,
    question: str,
    conversation_history: Optional[Sequence[Tuple[str, str]]] = None,
    filter_metadata: Optional[Dict[str, Any]] = None,
    doc_language: str = "unknown",
) -> Tuple[str, List[Dict[str, Any]]]:
    question_lang = resolve_response_language(question, doc_language)
    llm = _extract_llm_from_qa_chain(qa_chain)
    contextualized = contextualize_question(question, llm, conversation_history)
    retrieval_query = contextualized["standalone"]

    last_error: Exception | None = None
    docs: List[Document] = []
    for attempt in range(2):
        try:
            previous_filter = _set_retriever_filter(qa_chain.retriever, filter_metadata)
            try:
                docs, _used_queries = _retrieve_documents_cross_language(
                    qa_chain.retriever,
                    retrieval_query,
                    question_lang,
                    doc_language,
                    llm,
                )
            finally:
                _restore_retriever_filter(qa_chain.retriever, previous_filter)
            break
        except Exception as exc:
            last_error = exc
            if "llama runner process has terminated" not in str(exc) or attempt == 1:
                raise
            time.sleep(0.25)
    else:
        raise RuntimeError("Không thể sinh câu trả lời từ Ollama.") from last_error

    docs = _deduplicate_documents(list(docs or []))
    answer = _generate_answer_from_documents(
        llm,
        question,
        list(docs or []),
        question_lang,
        doc_language,
        conversation_history,
    )

    return answer, _format_source_documents(docs, len(docs))
